"""
Service de génération automatique de documents de soumission :
Offre Technique (PDF), Offre Financière / Bordereau des prix (XLSX),
Lettre de soumission et Déclaration d'engagement (DOCX).
"""
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

from sqlalchemy.orm import Session

from app.config import DATA_DIR
from app.models.soumission import Soumission, DocumentGenere
from app.models.appel_offres import AppelOffres

OUTPUT_DIR = DATA_DIR / "documents_generes"
OUTPUT_DIR.mkdir(exist_ok=True)

BLEU_INSTITUTIONNEL = RGBColor(0x0E, 0x3A, 0x5C)
OR_ACCENT = "C9912A"


def _generer_offre_technique_pdf(soumission: Soumission, appel_offres: AppelOffres | None) -> Path:
    filename = f"offre_technique_{soumission.id}.pdf"
    filepath = OUTPUT_DIR / filename

    reference_affichee = appel_offres.reference if appel_offres else (soumission.source_externe or "AO externe")
    objet_affiche = appel_offres.objet if appel_offres else soumission.objet

    doc = SimpleDocTemplate(str(filepath), pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm)
    styles = getSampleStyleSheet()
    titre_style = ParagraphStyle("Titre", parent=styles["Title"], textColor=colors.HexColor("#0E3A5C"), fontSize=18, spaceAfter=12)
    sous_titre_style = ParagraphStyle("SousTitre", parent=styles["Heading2"], textColor=colors.HexColor("#C9912A"), fontSize=13, spaceBefore=14, spaceAfter=6)
    corps_style = ParagraphStyle("Corps", parent=styles["Normal"], fontSize=10.5, leading=15)

    elements = []
    elements.append(Paragraph("JUMBO STORE SARL", titre_style))
    elements.append(Paragraph("OFFRE TECHNIQUE", styles["Heading1"]))
    elements.append(Paragraph(f"Référence de l'appel d'offres : {reference_affichee}", corps_style))
    elements.append(Spacer(1, 12))

    elements.append(Paragraph("1. Présentation de l'entreprise", sous_titre_style))
    elements.append(Paragraph(
        "Jumbo Store SARL est une entreprise ivoirienne forte de plus de 20 ans d'expérience dans le BTP "
        "(Bâtiment et Travaux Publics) et la vente de matériel général, basée à Cocody Attoban, Abidjan. "
        "L'entreprise dispose d'une expertise reconnue dans la réalisation de marchés publics et privés.",
        corps_style,
    ))

    elements.append(Paragraph("2. Objet du marché", sous_titre_style))
    elements.append(Paragraph(objet_affiche, corps_style))

    elements.append(Paragraph("3. Méthodologie proposée", sous_titre_style))
    methodo = soumission.methodologie_resume or (
        "Notre approche méthodologique s'appuie sur une planification rigoureuse en phases : "
        "(1) mobilisation des ressources humaines et matérielles, (2) exécution conforme au cahier des charges "
        "avec contrôle qualité continu, (3) livraison et réception définitive avec service après-vente."
    )
    elements.append(Paragraph(methodo, corps_style))

    elements.append(Paragraph("4. Planning d'exécution", sous_titre_style))
    planning = soumission.planning_resume or (
        "Le planning prévisionnel d'exécution sera communiqué en annexe, avec un délai global "
        "compatible avec les exigences du dossier d'appel d'offres."
    )
    elements.append(Paragraph(planning, corps_style))

    elements.append(Paragraph("5. Références techniques", sous_titre_style))
    elements.append(Paragraph(
        "Jumbo Store SARL a réalisé avec succès plusieurs marchés similaires pour des clients publics et privés "
        "en Côte d'Ivoire au cours des dernières années (détail en annexe sur demande).",
        corps_style,
    ))

    elements.append(Spacer(1, 24))
    elements.append(Paragraph(f"Fait à Abidjan, le {datetime.now().strftime('%d/%m/%Y')}", corps_style))
    elements.append(Paragraph("Pour Jumbo Store SARL", corps_style))

    doc.build(elements)
    return filepath


def _generer_offre_financiere_xlsx(soumission: Soumission, appel_offres: AppelOffres | None) -> Path:
    filename = f"offre_financiere_{soumission.id}.xlsx"
    filepath = OUTPUT_DIR / filename

    reference_affichee = appel_offres.reference if appel_offres else (soumission.source_externe or "AO externe")
    objet_affiche = appel_offres.objet if appel_offres else soumission.objet

    wb = Workbook()
    ws = wb.active
    ws.title = "Bordereau des Prix"

    header_fill = PatternFill(start_color="0E3A5C", end_color="0E3A5C", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True, size=11)
    title_font = Font(bold=True, size=14, color="0E3A5C")

    ws.merge_cells("A1:E1")
    ws["A1"] = "JUMBO STORE SARL — BORDEREAU DES PRIX UNITAIRES (BPU)"
    ws["A1"].font = title_font

    ws.merge_cells("A2:E2")
    ws["A2"] = f"Référence appel d'offres : {reference_affichee}"
    ws["A2"].font = Font(italic=True, size=10)

    headers = ["N°", "Désignation", "Quantité", "Prix unitaire (FCFA)", "Montant total (FCFA)"]
    for col, h in enumerate(headers, start=1):
        cell = ws.cell(row=4, column=col, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")

    montant_total = soumission.montant_propose or 0
    lignes_demo = [
        (1, f"Exécution conforme à l'objet : {objet_affiche[:60]}", 1, montant_total),
    ]
    row = 5
    for num, designation, qte, prix_unit in lignes_demo:
        ws.cell(row=row, column=1, value=num)
        ws.cell(row=row, column=2, value=designation)
        ws.cell(row=row, column=3, value=qte)
        ws.cell(row=row, column=4, value=prix_unit)
        ws.cell(row=row, column=5, value=qte * prix_unit)
        row += 1

    ws.cell(row=row + 1, column=4, value="MONTANT TOTAL HT").font = Font(bold=True)
    ws.cell(row=row + 1, column=5, value=montant_total).font = Font(bold=True)

    ws.column_dimensions["B"].width = 55
    ws.column_dimensions["A"].width = 6
    ws.column_dimensions["C"].width = 12
    ws.column_dimensions["D"].width = 20
    ws.column_dimensions["E"].width = 20

    wb.save(str(filepath))
    return filepath


def _generer_lettre_soumission_docx(soumission: Soumission, appel_offres: AppelOffres | None) -> Path:
    filename = f"lettre_soumission_{soumission.id}.docx"
    filepath = OUTPUT_DIR / filename

    reference_affichee = appel_offres.reference if appel_offres else (soumission.source_externe or "AO externe")
    objet_affiche = appel_offres.objet if appel_offres else soumission.objet

    doc = Document()

    titre = doc.add_paragraph()
    titre_run = titre.add_run("JUMBO STORE SARL")
    titre_run.bold = True
    titre_run.font.size = Pt(16)
    titre_run.font.color.rgb = BLEU_INSTITUTIONNEL
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Cocody Attoban, Abidjan — Côte d'Ivoire").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    sous_titre = doc.add_paragraph()
    sous_titre_run = sous_titre.add_run("LETTRE DE SOUMISSION")
    sous_titre_run.bold = True
    sous_titre_run.font.size = Pt(13)
    sous_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_paragraph(f"Référence de l'appel d'offres : {reference_affichee}")
    doc.add_paragraph(f"Objet : {objet_affiche}")
    doc.add_paragraph()

    corps = (
        "Monsieur/Madame le Président de la Commission,\n\n"
        "J'ai l'honneur de soumettre, par la présente, l'offre de Jumbo Store SARL pour l'exécution du marché "
        "cité en référence, conformément aux clauses et conditions du Dossier d'Appel d'Offres.\n\n"
        "Je déclare avoir pris connaissance de l'ensemble des pièces du dossier et m'engage, si mon offre est "
        "retenue, à exécuter le marché dans les conditions, délais et prix indiqués dans notre offre financière "
        "jointe.\n\n"
        "Je reste à votre disposition pour tout complément d'information."
    )
    doc.add_paragraph(corps)
    doc.add_paragraph()
    doc.add_paragraph(f"Fait à Abidjan, le {datetime.now().strftime('%d/%m/%Y')}")
    doc.add_paragraph()
    signature = doc.add_paragraph()
    signature.add_run("Pour Jumbo Store SARL").bold = True

    doc.save(str(filepath))
    return filepath


def generer_document(db: Session, soumission: Soumission, appel_offres: AppelOffres | None, type_document: str) -> DocumentGenere:
    if type_document == "Offre Technique":
        filepath = _generer_offre_technique_pdf(soumission, appel_offres)
        format_fichier = "PDF"
    elif type_document == "Offre Financière":
        filepath = _generer_offre_financiere_xlsx(soumission, appel_offres)
        format_fichier = "XLSX"
    elif type_document in ("Lettre de soumission", "Déclaration d'engagement"):
        filepath = _generer_lettre_soumission_docx(soumission, appel_offres)
        format_fichier = "DOCX"
    else:
        raise ValueError(f"Type de document non supporté : {type_document}")

    document = DocumentGenere(
        soumission_id=soumission.id,
        type_document=type_document,
        format_fichier=format_fichier,
        nom_fichier=filepath.name,
        chemin_fichier=str(filepath),
        date_generation=datetime.utcnow(),
    )
    db.add(document)
    db.commit()
    db.refresh(document)
    return document
